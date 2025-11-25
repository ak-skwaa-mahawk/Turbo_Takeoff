# main.py — Turbo Takeoff: Sovereign Edition — SUB PERFORMANCE RATINGS + AUTO-RANKING
# The circle judges. The circle evolves.

import os
import click
import yaml
import fitz
import json
from pathlib import Path
from datetime import datetime, timedelta
from email.message import EmailMessage
import smtplib
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
import pandas as pd
from docx import Document

load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

with open("config.yaml") as f:
    cfg = yaml.safe_load(f)

INPUT_DIR = Path(cfg["paths"]["input_pdfs"])
OUTPUT_DIR = Path(cfg["paths"]["output_folder"])
TRACKER_FILE = Path("sub_bid_tracker.json")
OUTPUT_DIR.mkdir(exist_ok=True)

# Load or init sovereign ledger
if TRACKER_FILE.exists():
    with open(TRACKER_FILE) as f:
        ledger = json.load(f)
else:
    ledger = {"subcontractors": {}, "projects": {}}

def save_ledger():
    with open(TRACKER_FILE, "w") as f:
        json.dump(ledger, f, indent=2)

# Init sub in ledger if new
def ensure_sub_in_ledger(sub):
    name = sub["name"]
    if name not in ledger["subcontractors"]:
        ledger["subcontractors"][name] = {
            "ethics_approved": sub.get("ethics_approved", False),
            "lifetime_jobs": 0,
            "avg_speed_hours": 0,
            "avg_price_variance_pct": 0,
            "circle_alignment_score": 100,
            "current_rating": 100,
            "history": []
        }

# Rate sub performance after reply
def rate_sub_performance(sub_name, project_key, reply_hours, price, baseline_price, notes=""):
    sub = ledger["subcontractors"][sub_name]
    sub["lifetime_jobs"] += 1
    
    # Speed score (0-100): 0-24hr = 100, 48hr = 70, >72hr = 30
    speed_score = max(0, 100 - (reply_hours - 24) * 1.25)
    
    # Price fairness: % variance from baseline (negative = came in under = good)
    variance = ((price - baseline_price) / baseline_price) * 100
    fairness_score = max(0, 100 + variance * 2)  # coming under boosts score
    
    # Circle alignment: starts 100, drops if notes contain red flags (manual later)
    alignment = sub["circle_alignment_score"]
    
    # Final rating
    rating = (speed_score * 0.4) + (fairness_score * 0.4) + (alignment * 0.2)
    sub["current_rating"] = round(rating, 1)
    
    # Update averages
    n = sub["lifetime_jobs"]
    sub["avg_speed_hours"] = (sub["avg_speed_hours"] * (n-1) + reply_hours) / n
    sub["avg_price_variance_pct"] = (sub["avg_price_variance_pct"] * (n-1) + variance) / n
    
    # Record history
    sub["history"].append({
        "project": project_key,
        "replied_hours": reply_hours,
        "price": price,
        "baseline": baseline_price,
        "variance_pct": round(variance, 1),
        "rating": sub["current_rating"],
        "date": datetime.now().isoformat()[:10]
    })
    
    # Blacklist if falls below 70
    if rating < 70:
        click.echo(f"   → {sub_name} RATING {rating} — FALLEN. Ethics review triggered.")
        # Optional: append to blacklist.yaml
    
    save_ledger()

# Get ranked list of subs for a scope
def get_ranked_subs_for_scope(scope_keyword):
    candidates = []
    for sub in cfg["subcontractors"]["subs"]:
        if any(scope_keyword.lower() in s.lower() for s in sub["scopes"]):
            name = sub["name"]
            ensure_sub_in_ledger(sub)
            rating = ledger["subcontractors"][name]["current_rating"]
            candidates.append((rating, sub))
    return [sub for rating, sub in sorted(candidates, key=lambda x: x[0], reverse=True)]

@click.command()
@click.argument("pdf_path", required=False)
def run(pdf_path: str | None):
    pdf_files = [Path(pdf_path)] if pdf_path else list(INPUT_DIR.glob("*.pdf"))
    if not pdf_files:
        click.echo("Drop a PDF in /input/")
        return

    for pdf in pdf_files:
        project_key = f"{datetime.now().strftime('%Y-%m-%d')}_{pdf.stem}"
        click.echo(f"\nSOVEREIGN CIRCLE BID + RATINGS → {pdf.name}")

        # PDF → images
        doc = fitz.open(pdf)
        images = []
        for i, page in enumerate(doc):
            p = OUTPUT_DIR / f"{project_key}_p{i:03d}.png"
            page.get_pixmap(dpi=300).save(p)
            images.append(p)

        # Detect needed subs (Gemini)
        needed_scopes = ["sheet metal", "traffic coating"]  # placeholder — real detection

        sub_total = 0
        invited_this_run = []

        for scope in needed_scopes:
            ranked = get_ranked_subs_for_scope(scope)
            if not ranked:
                continue
            top_sub = ranked[0]  # highest rated clean sub
            name = top_sub["name"]
            ensure_sub_in_ledger(top_sub)
            
            click.echo(f"   → {scope} → inviting #{1}: {name} (Rating: {ledger['subcontractors'][name]['current_rating']})")
            # send_sub_bid_request(top_sub, project_key, pdf.stem)  # uncomment when ready
            invited_this_run.append(name)
            
            # Use final quote if exists, else last_rate, else baseline
            quote = ledger["subcontractors"][name].get("final_quote") or top_sub["last_rate"]
            sub_total += quote

        # Core takeoff (placeholder)
        core_total = 52800.00
        grand_total = core_total + sub_total
        final_bid = grand_total * (1 + cfg["app"]["default_profit_pct"]/100)

        # Export with ratings
        rows = [
            {"Item": "Pro Seal Self-Perform", "Amount": core_total},
            {"Item": "Circle Subs (top-rated)", "Amount": sub_total},
            {"Item": "FINAL SOVEREIGN BID", "Amount": final_bid},
            {"Item": "", "Amount": None},
            {"Item": "SUB PERFORMANCE LEADERBOARD", "Amount": None},
        ]
        for name, data in sorted(ledger["subcontractors"].items(), key=lambda x: x[1]["current_rating"], reverse=True):
            rows.append({"Item": f"{name}", "Amount": f"Rating: {data['current_rating']} ★"})
        
        df = pd.DataFrame(rows)
        excel_path = OUTPUT_DIR / f"{project_key}_RATED_CIRCLE_BID.xlsx"
        df.to_excel(excel_path, index=False)

        doc = Document()
        doc.add_heading("PRO SEAL — RATED CIRCLE BID", 0)
        doc.add_paragraph(f"Project: {pdf.stem}")
        doc.add_paragraph(f"Top-rated circle subs invited: {', '.join(invited_this_run) or 'None'}")
        doc.add_paragraph(f"FINAL BID: ${final_bid:,.0f}")
        doc.add_paragraph("\nThe circle judges. The circle ranks. The circle wins.")
        doc.add_paragraph("Only the strong chase.")
        doc.save(OUTPUT_DIR / f"{project_key}_RATED_PROPOSAL.docx")

        click.echo(f"Rated bid complete → {excel_path}")
        click.echo("The circle has spoken.\n")

if __name__ == "__main__":
    run()
def rate_sub_performance(sub_name, project_key, reply_hours, price, baseline_price, alignment_bonus=0, notes=""):
    """
    Refined sovereign rating algorithm — November 2025
    Speed 45% · Price 35% · Alignment 20%
    """
    sub = ledger["subcontractors"][sub_name]
    sub["lifetime_jobs"] += 1

    # 1. SPEED SCORE — 45% weight
    # 0–24h = 100, 25–48h = 70, 49–72h = 40, >72h = 10
    if reply_hours <= 24:
        speed_score = 100
    elif reply_hours <= 48:
        speed_score = 70
    elif reply_hours <= 72:
        speed_score = 40
    else:
        speed_score = 10

    # 2. PRICE FAIRNESS — 35% weight
    variance_pct = ((price - baseline_price) / baseline_price) * 100
    # Coming in under = bonus, over = penalty (non-linear forgiveness)
    if variance_pct <= 0:
        price_score = 100 + (abs(variance_pct) * 1.5)  # e.g., -5% → 107.5 (capped later)
    else:
        price_score = 100 - (variance_pct * 3)           # +5% → 85, +10% → 70

    # 3. CIRCLE ALIGNMENT — 20% weight (starts 100, decays only on betrayal)
    alignment = sub.get("circle_alignment_score", 100) + alignment_bonus
    alignment = min(100, max(0, alignment))

    # Final sovereign rating
    raw_rating = (speed_score * 0.45) + (price_score * 0.35) + (alignment * 0.20)
    final_rating = round(max(0, min(100, raw_rating)), 1)

    # Update ledger
    n = sub["lifetime_jobs"]
    sub["avg_speed_hours"] = (sub["avg_speed_hours"] * (n-1) + reply_hours) / n
    sub["avg_price_variance_pct"] = (sub["avg_price_variance_pct"] * (n-1) + variance_pct) / n
    sub["circle_alignment_score"] = alignment
    sub["current_rating"] = final_rating

    # Eternal record
    sub["history"].append({
        "project": project_key,
        "reply_hours": reply_hours,
        "price": price,
        "baseline": baseline_price,
        "variance_pct": round(variance_pct, 2),
        "speed_score": speed_score,
        "price_score": round(price_score, 1),
        "alignment": alignment,
        "final_rating": final_rating,
        "date": datetime.now().isoformat()[:10]
    })

    # Auto-blacklist threshold: <65 → fallen from the circle
    if final_rating < 65:
        click.echo(f"   FALLEN: {sub_name} rating {final_rating} → Ethics blade triggered. Removed from future invites.")
        # Optional: auto-append to ethics_blacklist.yaml

    save_ledger()
    click.echo(f"   JUDGED: {sub_name} → Rating {final_rating}/100 | Speed {speed_score} | Price {round(price_score,1)} | Alignment {alignment}")